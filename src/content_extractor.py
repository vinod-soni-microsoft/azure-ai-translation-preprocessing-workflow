from docx import Document
from typing import Dict, List, Optional, Tuple
import re
import logging
from pathlib import Path
from .azure_translate_content_analyzer import AzureTranslateContentAnalyzer

logger = logging.getLogger(__name__)

class ContentExtractor:
    """Extracts and analyzes content from DOCX documents with Azure AI Translate compatibility."""
    
    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.azure_analyzer = AzureTranslateContentAnalyzer()
    
    def extract_content(self, file_path: str) -> Dict:
        """
        Extracts all text content from a DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract paragraph text
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract table content
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        table_data.append(row_data)
                if table_data:
                    tables_content.append(table_data)
            
            # Extract headers and footers
            headers_footers = []
            for section in doc.sections:
                # Headers
                if section.header.paragraphs:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            headers_footers.append(f"Header: {para.text.strip()}")
                
                # Footers
                if section.footer.paragraphs:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            headers_footers.append(f"Footer: {para.text.strip()}")
            
            # Combine all text
            all_text = []
            all_text.extend(paragraphs)
            
            # Flatten table content
            for table in tables_content:
                for row in table:
                    all_text.extend(row)
            
            all_text.extend(headers_footers)
            
            full_text = " ".join(all_text)
            
            result = {
                'paragraphs': paragraphs,
                'tables': tables_content,
                'headers_footers': headers_footers,
                'full_text': full_text,
                'content_analysis': self._analyze_content(full_text, paragraphs, tables_content)
            }
            
            self.logger.info(f"Content extracted successfully from {file_path}")
            return result
            
        except Exception as e:
            error_msg = f"Error extracting content from {file_path}: {str(e)}"
            self.logger.error(error_msg)
            return {'error': error_msg}
    
    def _analyze_content(self, full_text: str, paragraphs: List[str], tables: List) -> Dict:
        """
        Analyzes the extracted content for translatability and other metrics.
        
        Args:
            full_text: Combined text from the document
            paragraphs: List of paragraph texts
            tables: List of table data
            
        Returns:
            Dictionary with content analysis results
        """
        try:
            analysis = {
                'has_translatable_text': False,
                'total_characters': len(full_text),
                'total_words': len(full_text.split()) if full_text else 0,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'estimated_translatable_words': 0,
                'language_hints': [],
                'content_types': []
            }
            
            # Check for translatable text (non-empty, non-numeric content)
            if full_text.strip():
                # Remove pure numbers, dates, and common non-translatable patterns
                translatable_text = re.sub(r'\b\d+\b', '', full_text)  # Remove numbers
                translatable_text = re.sub(r'\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}\b', '', translatable_text)  # Remove dates
                translatable_text = re.sub(r'[^\w\s]', ' ', translatable_text)  # Remove special chars
                translatable_text = ' '.join(translatable_text.split())  # Clean whitespace
                
                if len(translatable_text.strip()) > 10:  # Minimum threshold for translatable content
                    analysis['has_translatable_text'] = True
                    analysis['estimated_translatable_words'] = len(translatable_text.split())
            
            # Detect content types
            if paragraphs:
                analysis['content_types'].append('paragraphs')
            if tables:
                analysis['content_types'].append('tables')
            
            # Simple language detection hints
            if re.search(r'[a-zA-Z]', full_text):
                analysis['language_hints'].append('contains_latin_script')
            if re.search(r'[àáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿ]', full_text.lower()):
                analysis['language_hints'].append('contains_accented_characters')
            
            return analysis
            
        except Exception as e:
            self.logger.error(f"Error analyzing content: {e}")
            return {'error': str(e)}
    
    def check_translatable_content(self, file_path: str) -> Tuple[bool, str, Dict]:
        """
        Checks if a document contains translatable content using Azure AI Translate standards.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (has_content, message, detailed_analysis)
        """
        try:
            # Use Azure AI Translate analyzer for comprehensive analysis
            azure_analysis = self.azure_analyzer.analyze_for_azure_translate(file_path)
            
            if 'error' in azure_analysis:
                return False, f"Failed to analyze content: {azure_analysis['error']}", {}
            
            # Get legacy analysis for backward compatibility
            content_data = self.extract_content(file_path)
            if 'error' in content_data:
                legacy_analysis = {}
            else:
                legacy_analysis = content_data.get('content_analysis', {})
            
            # Determine if content is translatable based on Azure standards
            is_translatable = azure_analysis['azure_translate_ready']
            readiness_score = azure_analysis['readiness_score']
            translation_analysis = azure_analysis['translation_analysis']
            
            # Create comprehensive analysis combining both approaches
            detailed_analysis = {
                # Legacy fields for backward compatibility
                'has_translatable_text': is_translatable,
                'total_characters': translation_analysis['total_characters'],
                'total_words': translation_analysis['total_words'],
                'estimated_translatable_words': translation_analysis['translatable_words'],
                
                # Enhanced Azure AI Translate fields
                'azure_translate_ready': is_translatable,
                'readiness_score': readiness_score,
                'translatable_segments': translation_analysis['translatable_segments'],
                'content_types': translation_analysis['content_types'],
                'text_density': translation_analysis['text_density'],
                'translation_complexity': translation_analysis['translation_complexity'],
                
                # Language and compatibility info
                'language_analysis': azure_analysis['language_analysis'],
                'azure_compatibility': azure_analysis['azure_compatibility'],
                'recommendations': azure_analysis['recommendations'],
                
                # Segmentation info
                'segmentation_analysis': azure_analysis.get('segmentation_analysis', {}),
                'formatting_analysis': azure_analysis.get('formatting_analysis', {})
            }
            
            if is_translatable:
                message = (
                    f"Document is ready for Azure AI Translate "
                    f"({translation_analysis['translatable_words']} translatable words, "
                    f"{readiness_score:.1%} readiness score)"
                )
            else:
                failed_requirements = azure_analysis['azure_compatibility']['requirements_failed']
                message = (
                    f"Document not ready for Azure AI Translate. "
                    f"Issues: {'; '.join(failed_requirements[:2])}"
                )
                if len(failed_requirements) > 2:
                    message += f" and {len(failed_requirements) - 2} more"
            
            return is_translatable, message, detailed_analysis
                
        except Exception as e:
            error_msg = f"Error checking translatable content: {str(e)}"
            self.logger.error(error_msg)
            return False, error_msg, {}
    
    def get_azure_translate_summary(self, file_path: str) -> Dict:
        """
        Get a concise Azure AI Translate readiness summary.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with Azure AI Translate readiness summary
        """
        try:
            return self.azure_analyzer.get_azure_translate_summary(file_path)
        except Exception as e:
            self.logger.error(f"Error getting Azure translate summary: {e}")
            return {'error': str(e)}
