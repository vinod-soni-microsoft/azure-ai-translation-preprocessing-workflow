"""
Enhanced content analyzer specifically designed for Azure AI Translate service compatibility.
Implements Azure AI Translate standards for document preprocessing and validation.
"""

from docx import Document
from typing import Dict, List, Optional, Tuple, Set
import re
import logging
from pathlib import Path
import unicodedata
from collections import defaultdict

logger = logging.getLogger(__name__)

class AzureTranslateContentAnalyzer:
    """
    Advanced content analyzer that follows Azure AI Translate service standards
    for optimal document translation preparation and validation.
    """
    
    # Azure AI Translate supported languages (subset for demonstration)
    SUPPORTED_LANGUAGES = {
        'en', 'es', 'fr', 'de', 'it', 'pt', 'ru', 'zh', 'ja', 'ko', 
        'ar', 'hi', 'tr', 'pl', 'nl', 'sv', 'da', 'no', 'fi'
    }
    
    # Minimum thresholds for Azure AI Translate
    MIN_TRANSLATABLE_CHARS = 3
    MIN_TRANSLATABLE_WORDS = 1
    MAX_SEGMENT_LENGTH = 5000  # Azure AI Translate segment limit
    
    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
        
    def analyze_for_azure_translate(self, file_path: str) -> Dict:
        """
        Comprehensive analysis for Azure AI Translate service compatibility.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with Azure AI Translate compatibility analysis
        """
        try:
            doc = Document(file_path)
            
            # Extract structured content
            content_structure = self._extract_structured_content(doc)
            
            # Analyze translatability
            translation_analysis = self._analyze_translation_readiness(content_structure)
            
            # Language detection
            language_analysis = self._detect_languages(content_structure['all_text'])
            
            # Segmentation analysis
            segmentation_analysis = self._analyze_segmentation(content_structure['segments'])
            
            # Formatting preservation analysis
            formatting_analysis = self._analyze_formatting_preservation(doc)
            
            # Azure AI specific validations
            azure_compatibility = self._validate_azure_ai_compatibility(
                content_structure, translation_analysis, language_analysis
            )
            
            result = {
                'azure_translate_ready': azure_compatibility['is_ready'],
                'readiness_score': azure_compatibility['score'],
                'content_structure': content_structure,
                'translation_analysis': translation_analysis,
                'language_analysis': language_analysis,
                'segmentation_analysis': segmentation_analysis,
                'formatting_analysis': formatting_analysis,
                'azure_compatibility': azure_compatibility,
                'recommendations': azure_compatibility['recommendations']
            }
            
            self.logger.info(f"Azure AI Translate analysis completed for {file_path}")
            return result
            
        except Exception as e:
            error_msg = f"Error in Azure AI Translate analysis for {file_path}: {str(e)}"
            self.logger.error(error_msg)
            return {'error': error_msg}
    
    def _extract_structured_content(self, doc: Document) -> Dict:
        """Extract content in structured format for translation analysis."""
        structure = {
            'paragraphs': [],
            'tables': [],
            'headers_footers': [],
            'text_runs': [],
            'segments': [],
            'all_text': '',
            'total_elements': 0
        }
        
        # Extract paragraphs with formatting info
        for para in doc.paragraphs:
            if para.text.strip():
                para_info = {
                    'text': para.text.strip(),
                    'style': para.style.name if para.style else 'Normal',
                    'has_formatting': any(run.bold or run.italic or run.underline for run in para.runs)
                }
                structure['paragraphs'].append(para_info)
                structure['segments'].append(para.text.strip())
        
        # Extract table content with structure
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'table_id': table_idx,
                'rows': [],
                'has_headers': False
            }
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                        structure['segments'].append(cell.text.strip())
                
                if row_data:
                    table_data['rows'].append(row_data)
                    # Assume first row might be headers
                    if row_idx == 0:
                        table_data['has_headers'] = True
            
            if table_data['rows']:
                structure['tables'].append(table_data)
        
        # Extract headers and footers
        for section in doc.sections:
            # Headers
            for para in section.header.paragraphs:
                if para.text.strip():
                    header_text = f"Header: {para.text.strip()}"
                    structure['headers_footers'].append(header_text)
                    structure['segments'].append(para.text.strip())
            
            # Footers  
            for para in section.footer.paragraphs:
                if para.text.strip():
                    footer_text = f"Footer: {para.text.strip()}"
                    structure['headers_footers'].append(footer_text)
                    structure['segments'].append(para.text.strip())
        
        # Combine all text
        all_texts = []
        all_texts.extend([p['text'] for p in structure['paragraphs']])
        
        for table in structure['tables']:
            for row in table['rows']:
                all_texts.extend(row)
        
        all_texts.extend(structure['headers_footers'])
        
        structure['all_text'] = ' '.join(all_texts)
        structure['total_elements'] = len(structure['segments'])
        
        return structure
    
    def _analyze_translation_readiness(self, content_structure: Dict) -> Dict:
        """Analyze content for translation readiness per Azure AI standards."""
        all_text = content_structure['all_text']
        segments = content_structure['segments']
        
        analysis = {
            'has_translatable_content': False,
            'translatable_segments': 0,
            'non_translatable_segments': 0,
            'total_characters': len(all_text),
            'total_words': len(all_text.split()) if all_text else 0,
            'translatable_characters': 0,
            'translatable_words': 0,
            'content_types': [],
            'text_density': 0.0,
            'translation_complexity': 'low'
        }
        
        translatable_text_parts = []
        
        for segment in segments:
            if self._is_segment_translatable(segment):
                analysis['translatable_segments'] += 1
                translatable_text_parts.append(segment)
            else:
                analysis['non_translatable_segments'] += 1
        
        if translatable_text_parts:
            combined_translatable = ' '.join(translatable_text_parts)
            analysis['translatable_characters'] = len(combined_translatable)
            analysis['translatable_words'] = len(combined_translatable.split())
            analysis['has_translatable_content'] = (
                analysis['translatable_characters'] >= self.MIN_TRANSLATABLE_CHARS and
                analysis['translatable_words'] >= self.MIN_TRANSLATABLE_WORDS
            )
        
        # Calculate text density
        if analysis['total_characters'] > 0:
            analysis['text_density'] = analysis['translatable_characters'] / analysis['total_characters']
        
        # Determine content types
        if content_structure['paragraphs']:
            analysis['content_types'].append('text_paragraphs')
        if content_structure['tables']:
            analysis['content_types'].append('structured_tables')
        if content_structure['headers_footers']:
            analysis['content_types'].append('headers_footers')
        
        # Assess translation complexity
        analysis['translation_complexity'] = self._assess_complexity(all_text)
        
        return analysis
    
    def _is_segment_translatable(self, text: str) -> bool:
        """Determine if a text segment is translatable per Azure AI standards."""
        if not text or len(text.strip()) < self.MIN_TRANSLATABLE_CHARS:
            return False
        
        # Remove common non-translatable patterns
        cleaned_text = text.strip()
        
        # Skip pure numbers
        if re.match(r'^\d+\.?\d*$', cleaned_text):
            return False
        
        # Skip pure dates
        if re.match(r'^\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}$', cleaned_text):
            return False
        
        # Skip URLs and emails
        if re.match(r'^https?://|^www\.|@.*\.[a-z]{2,}$', cleaned_text.lower()):
            return False
        
        # Skip pure punctuation or symbols
        if re.match(r'^[^\w\s]+$', cleaned_text):
            return False
        
        # Check for meaningful content (letters)
        if not re.search(r'[a-zA-ZàáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿA-ZÀ-Ÿ]', cleaned_text):
            return False
        
        # Must have at least one word character
        word_chars = re.findall(r'\w+', cleaned_text)
        if not word_chars or len(' '.join(word_chars)) < self.MIN_TRANSLATABLE_CHARS:
            return False
        
        return True
    
    def _detect_languages(self, text: str) -> Dict:
        """Simple language detection hints for Azure AI Translate."""
        analysis = {
            'detected_scripts': [],
            'likely_languages': [],
            'is_multilingual': False,
            'confidence': 'low',
            'azure_supported': True
        }
        
        if not text:
            return analysis
        
        # Detect scripts
        scripts = set()
        for char in text:
            script = unicodedata.name(char, '').split()[0] if unicodedata.name(char, '') else ''
            if script:
                scripts.add(script)
        
        # Basic script detection
        if any(re.search(pattern, text) for pattern in [r'[a-zA-Z]']):
            analysis['detected_scripts'].append('Latin')
            analysis['likely_languages'].extend(['en', 'es', 'fr', 'de', 'it'])
        
        if any(re.search(pattern, text) for pattern in [r'[àáâãäåæçèéêëìíîïðñòóôõöøùúûüýþÿ]']):
            analysis['detected_scripts'].append('Latin_Extended')
            analysis['likely_languages'].extend(['es', 'fr', 'de', 'pt'])
        
        if re.search(r'[\u4e00-\u9fff]', text):
            analysis['detected_scripts'].append('Chinese')
            analysis['likely_languages'].append('zh')
        
        if re.search(r'[\u3040-\u309f\u30a0-\u30ff]', text):
            analysis['detected_scripts'].append('Japanese')
            analysis['likely_languages'].append('ja')
        
        if re.search(r'[\u0600-\u06ff]', text):
            analysis['detected_scripts'].append('Arabic')
            analysis['likely_languages'].append('ar')
        
        # Remove duplicates and check Azure support
        analysis['likely_languages'] = list(set(analysis['likely_languages']))
        analysis['azure_supported'] = all(
            lang in self.SUPPORTED_LANGUAGES 
            for lang in analysis['likely_languages']
        )
        
        # Simple multilingual detection
        analysis['is_multilingual'] = len(analysis['detected_scripts']) > 1
        
        if analysis['likely_languages']:
            analysis['confidence'] = 'medium'
        
        return analysis
    
    def _analyze_segmentation(self, segments: List[str]) -> Dict:
        """Analyze text segmentation for Azure AI Translate optimization."""
        analysis = {
            'total_segments': len(segments),
            'segments_within_limit': 0,
            'segments_exceeding_limit': 0,
            'average_segment_length': 0,
            'max_segment_length': 0,
            'requires_segmentation': False,
            'optimal_for_translation': True
        }
        
        if not segments:
            return analysis
        
        segment_lengths = []
        
        for segment in segments:
            length = len(segment)
            segment_lengths.append(length)
            
            if length <= self.MAX_SEGMENT_LENGTH:
                analysis['segments_within_limit'] += 1
            else:
                analysis['segments_exceeding_limit'] += 1
        
        analysis['average_segment_length'] = sum(segment_lengths) / len(segment_lengths)
        analysis['max_segment_length'] = max(segment_lengths)
        analysis['requires_segmentation'] = analysis['segments_exceeding_limit'] > 0
        analysis['optimal_for_translation'] = (
            analysis['segments_exceeding_limit'] == 0 and
            analysis['average_segment_length'] < self.MAX_SEGMENT_LENGTH * 0.8
        )
        
        return analysis
    
    def _analyze_formatting_preservation(self, doc: Document) -> Dict:
        """Analyze formatting elements that Azure AI Translate can preserve."""
        analysis = {
            'has_formatting': False,
            'formatting_elements': [],
            'preservation_supported': True,
            'complex_formatting': False
        }
        
        formatting_found = set()
        
        # Check paragraph formatting
        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold:
                    formatting_found.add('bold')
                if run.italic:
                    formatting_found.add('italic')
                if run.underline:
                    formatting_found.add('underline')
                if run.font.size:
                    formatting_found.add('font_size')
                if run.font.color.rgb:
                    formatting_found.add('font_color')
        
        # Check for tables (structure preservation)
        if doc.tables:
            formatting_found.add('tables')
        
        # Check for headers/footers
        for section in doc.sections:
            if any(p.text.strip() for p in section.header.paragraphs):
                formatting_found.add('headers')
            if any(p.text.strip() for p in section.footer.paragraphs):
                formatting_found.add('footers')
        
        analysis['formatting_elements'] = list(formatting_found)
        analysis['has_formatting'] = len(formatting_found) > 0
        analysis['complex_formatting'] = len(formatting_found) > 3
        
        # Azure AI Translate supports most basic formatting
        analysis['preservation_supported'] = True
        
        return analysis
    
    def _assess_complexity(self, text: str) -> str:
        """Assess translation complexity based on content characteristics."""
        if not text:
            return 'none'
        
        complexity_score = 0
        
        # Check for technical terms (uppercase words)
        technical_terms = len(re.findall(r'\b[A-Z]{2,}\b', text))
        if technical_terms > 5:
            complexity_score += 1
        
        # Check for numbers and measurements
        numbers = len(re.findall(r'\d+', text))
        if numbers > 10:
            complexity_score += 1
        
        # Check for special characters
        special_chars = len(re.findall(r'[^\w\s]', text))
        if special_chars > len(text) * 0.1:
            complexity_score += 1
        
        # Check sentence length
        sentences = re.split(r'[.!?]+', text)
        avg_sentence_length = sum(len(s.split()) for s in sentences if s.strip()) / max(len(sentences), 1)
        if avg_sentence_length > 25:
            complexity_score += 1
        
        if complexity_score == 0:
            return 'low'
        elif complexity_score <= 2:
            return 'medium'
        else:
            return 'high'
    
    def _validate_azure_ai_compatibility(self, content_structure: Dict, 
                                       translation_analysis: Dict, 
                                       language_analysis: Dict) -> Dict:
        """Validate overall compatibility with Azure AI Translate service."""
        compatibility = {
            'is_ready': False,
            'score': 0.0,
            'requirements_met': [],
            'requirements_failed': [],
            'recommendations': []
        }
        
        total_checks = 6
        passed_checks = 0
        
        # Check 1: Has translatable content
        if translation_analysis['has_translatable_content']:
            compatibility['requirements_met'].append('Has sufficient translatable content')
            passed_checks += 1
        else:
            compatibility['requirements_failed'].append('Insufficient translatable content')
            compatibility['recommendations'].append(
                'Add more meaningful text content (minimum 3 characters, 1 word)'
            )
        
        # Check 2: Content density
        if translation_analysis['text_density'] > 0.1:  # At least 10% translatable
            compatibility['requirements_met'].append('Good text density')
            passed_checks += 1
        else:
            compatibility['requirements_failed'].append('Low text density')
            compatibility['recommendations'].append(
                'Increase ratio of translatable text to total content'
            )
        
        # Check 3: Segment size optimization
        segmentation = content_structure.get('segmentation_analysis', {})
        if not segmentation.get('requires_segmentation', True):
            compatibility['requirements_met'].append('Optimal segment sizes')
            passed_checks += 1
        else:
            compatibility['requirements_failed'].append('Segments too large')
            compatibility['recommendations'].append(
                f'Break large segments into smaller parts (max {self.MAX_SEGMENT_LENGTH} chars)'
            )
        
        # Check 4: Language support
        if language_analysis['azure_supported']:
            compatibility['requirements_met'].append('Supported languages detected')
            passed_checks += 1
        else:
            compatibility['requirements_failed'].append('Unsupported language detected')
            compatibility['recommendations'].append(
                'Verify document language is supported by Azure AI Translate'
            )
        
        # Check 5: Content structure
        if content_structure['total_elements'] > 0:
            compatibility['requirements_met'].append('Valid content structure')
            passed_checks += 1
        else:
            compatibility['requirements_failed'].append('No structured content found')
            compatibility['recommendations'].append(
                'Ensure document has proper paragraph or table structure'
            )
        
        # Check 6: Format preservation
        if translation_analysis['total_words'] > 0:
            compatibility['requirements_met'].append('Format suitable for preservation')
            passed_checks += 1
        else:
            compatibility['requirements_failed'].append('Format may not preserve well')
            compatibility['recommendations'].append(
                'Verify document formatting is compatible with DOCX preservation'
            )
        
        # Calculate score and readiness
        compatibility['score'] = passed_checks / total_checks
        compatibility['is_ready'] = compatibility['score'] >= 0.8  # 80% threshold
        
        # Add general recommendations
        if compatibility['is_ready']:
            compatibility['recommendations'].append(
                'Document is ready for Azure AI Translate service'
            )
        else:
            compatibility['recommendations'].append(
                f'Address {len(compatibility["requirements_failed"])} issues before translation'
            )
        
        return compatibility
    
    def get_azure_translate_summary(self, file_path: str) -> Dict:
        """Get a concise summary for Azure AI Translate readiness."""
        analysis = self.analyze_for_azure_translate(file_path)
        
        if 'error' in analysis:
            return analysis
        
        return {
            'ready_for_translation': analysis['azure_translate_ready'],
            'readiness_score': f"{analysis['readiness_score']:.1%}",
            'translatable_words': analysis['translation_analysis']['translatable_words'],
            'detected_languages': analysis['language_analysis']['likely_languages'],
            'content_types': analysis['translation_analysis']['content_types'],
            'key_recommendations': analysis['recommendations'][:3],  # Top 3 recommendations
            'azure_compatibility': analysis['azure_compatibility']['is_ready']
        }
