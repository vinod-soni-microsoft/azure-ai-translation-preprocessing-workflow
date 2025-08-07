#!/usr/bin/env python3
"""
Quick start script for the Document Processing Service
This script helps users get started quickly with minimal setup
"""

import os
import sys
import subprocess
import platform
from pathlib import Path

def print_header():
    """Print welcome header."""
    print("🚀 Document Processing Service - Quick Start")
    print("=" * 50)

def check_python():
    """Check Python version."""
    print("\n📋 Checking Python installation...")
    
    version = sys.version_info
    if version.major >= 3 and version.minor >= 8:
        print(f"✅ Python {version.major}.{version.minor}.{version.micro} - OK")
        return True
    else:
        print(f"❌ Python {version.major}.{version.minor}.{version.micro} - Need Python 3.8+")
        return False

def check_dependencies():
    """Check if dependencies are installed."""
    print("\n📦 Checking dependencies...")
    
    try:
        import fastapi
        import uvicorn
        import docx
        print("✅ All dependencies installed")
        return True
    except ImportError as e:
        print(f"❌ Missing dependencies: {e}")
        return False

def install_dependencies():
    """Install dependencies."""
    print("\n📥 Installing dependencies...")
    
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-r", "requirements.txt"])
        print("✅ Dependencies installed successfully")
        return True
    except subprocess.CalledProcessError:
        print("❌ Failed to install dependencies")
        return False

def check_libreoffice():
    """Check LibreOffice installation."""
    print("\n📄 Checking LibreOffice installation...")
    
    # Common LibreOffice paths
    common_paths = {
        "Windows": [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
        ],
        "Darwin": [  # macOS
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        ],
        "Linux": [
            "/usr/bin/libreoffice",
            "/usr/local/bin/libreoffice"
        ]
    }
    
    system = platform.system()
    paths = common_paths.get(system, [])
    
    # Check if in PATH
    try:
        subprocess.run(["soffice", "--version"], capture_output=True, check=True)
        print("✅ LibreOffice found in PATH")
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        pass
    
    # Check common installation paths
    for path in paths:
        if os.path.exists(path):
            print(f"✅ LibreOffice found at: {path}")
            return True
    
    print("⚠️  LibreOffice not found - Format conversion will be limited")
    print("   Install LibreOffice for full functionality:")
    if system == "Windows":
        print("   - Download from: https://www.libreoffice.org/download/")
    elif system == "Darwin":
        print("   - Run: brew install libreoffice")
    elif system == "Linux":
        print("   - Run: sudo apt-get install libreoffice")
    
    return False

def create_directories():
    """Create necessary directories."""
    print("\n📁 Creating directories...")
    
    directories = ["uploads", "converted", "logs"]
    
    for directory in directories:
        Path(directory).mkdir(exist_ok=True)
        print(f"✅ Created: {directory}/")

def create_sample_files():
    """Create sample test files."""
    print("\n📝 Creating sample files...")
    
    try:
        from docx import Document
        
        # Create a sample DOCX file
        doc = Document()
        doc.add_heading('Sample Document for Testing', 0)
        doc.add_paragraph('This is a sample document that contains translatable text content.')
        doc.add_paragraph('It includes multiple paragraphs with meaningful content that should be detected by the content extractor.')
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'
        
        doc.add_paragraph('This document should pass all validation checks.')
        doc.save('uploads/sample_valid.docx')
        print("✅ Created: uploads/sample_valid.docx")
        
        # Create test files for negative scenarios
        with open('uploads/fake_docx.docx', 'w') as f:
            f.write('This is not a real DOCX file, just text with .docx extension')
        print("✅ Created: uploads/fake_docx.docx (for testing)")
        
        with open('uploads/sample.txt', 'w') as f:
            f.write('This is a plain text file that needs conversion to DOCX format.')
        print("✅ Created: uploads/sample.txt (for conversion testing)")
        
        return True
        
    except ImportError:
        print("❌ Cannot create sample DOCX - python-docx not installed")
        return False
    except Exception as e:
        print(f"❌ Error creating sample files: {e}")
        return False

def start_service():
    """Start the FastAPI service."""
    print("\n🚀 Starting the Document Processing Service...")
    print("   Service will be available at: http://localhost:8000")
    print("   API Documentation: http://localhost:8000/docs")
    print("   Press Ctrl+C to stop the service")
    print("\n" + "-" * 50)
    
    try:
        os.chdir("src")
        subprocess.run([sys.executable, "main.py"])
    except KeyboardInterrupt:
        print("\n\n👋 Service stopped by user")
    except Exception as e:
        print(f"\n❌ Error starting service: {e}")

def run_tests():
    """Run basic tests."""
    print("\n🧪 Running basic tests...")
    
    try:
        result = subprocess.run([sys.executable, "test_service.py"], capture_output=True, text=True)
        if result.returncode == 0:
            print("✅ Basic tests passed")
            return True
        else:
            print(f"❌ Tests failed: {result.stderr}")
            return False
    except Exception as e:
        print(f"❌ Error running tests: {e}")
        return False

def main_menu():
    """Display main menu and handle user choices."""
    while True:
        print("\n🎯 Quick Start Menu:")
        print("1. 🔧 Setup (install dependencies, create directories)")
        print("2. 📝 Create sample files")
        print("3. 🧪 Run tests")
        print("4. 🚀 Start service")
        print("5. ❓ Show service info")
        print("6. 🚪 Exit")
        
        choice = input("\nSelect an option (1-6): ").strip()
        
        if choice == "1":
            setup()
        elif choice == "2":
            create_sample_files()
        elif choice == "3":
            run_tests()
        elif choice == "4":
            start_service()
        elif choice == "5":
            show_service_info()
        elif choice == "6":
            print("👋 Goodbye!")
            break
        else:
            print("❌ Invalid choice. Please select 1-6.")

def setup():
    """Run complete setup."""
    print("\n🔧 Running setup...")
    
    # Check Python
    if not check_python():
        print("❌ Setup failed - Python version too old")
        return
    
    # Install dependencies
    if not check_dependencies():
        if not install_dependencies():
            print("❌ Setup failed - Could not install dependencies")
            return
    
    # Check LibreOffice
    check_libreoffice()
    
    # Create directories
    create_directories()
    
    print("\n✅ Setup completed successfully!")
    print("   Next steps:")
    print("   1. Create sample files (option 2)")
    print("   2. Run tests (option 3)")
    print("   3. Start service (option 4)")

def show_service_info():
    """Show service information."""
    print("\n📋 Document Processing Service Information:")
    print("   📖 Full Documentation: COMPLETE_GUIDE.md")
    print("   🧪 Test Results: TEST_RESULTS.md")
    print("   📝 README: README.md")
    print("\n🌐 When running, the service provides:")
    print("   • REST API at: http://localhost:8000")
    print("   • Interactive docs: http://localhost:8000/docs")
    print("   • Health check: http://localhost:8000/health")
    print("\n📁 Directory structure:")
    print("   • uploads/ - For uploaded files")
    print("   • converted/ - For converted files")
    print("   • logs/ - For operation logs")
    print("   • src/ - Source code")

def main():
    """Main function."""
    print_header()
    
    # Quick system check
    print("🔍 Quick system check...")
    python_ok = check_python()
    deps_ok = check_dependencies()
    
    if python_ok and deps_ok:
        print("✅ System ready!")
    else:
        print("⚠️  System needs setup")
    
    main_menu()

if __name__ == "__main__":
    main()
